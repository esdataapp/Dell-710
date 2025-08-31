#!/usr/bin/env python3
"""
Markdown → DOCX simple converter for PropertyScraper docs
- Minimal, dependency-light (uses python-docx only)
- Supports: headings (#..######), paragraphs, bullet/numbered lists, code blocks
"""

from __future__ import annotations
import argparse
from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_heading(doc: Document, text: str, level: int) -> None:
    level = max(1, min(level, 4))  # limit to 1-4
    doc.add_heading(text.strip(), level=level)


def add_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text.rstrip())
    p_format = p.paragraph_format
    p_format.space_after = Pt(6)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text.strip())
    p.style = 'List Bullet'


def add_numbered(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text.strip())
    p.style = 'List Number'


def add_code(doc: Document, code_lines: list[str]) -> None:
    for line in code_lines:
        p = doc.add_paragraph()
        run = p.add_run(line.rstrip('\n'))
        run.font.name = 'Consolas'
        run.font.size = Pt(9)


def convert_markdown_to_docx(md_path: Path, out_path: Path, title: str | None = None) -> None:
    text = md_path.read_text(encoding='utf-8', errors='replace')
    lines = text.splitlines()

    doc = Document()

    # Title
    doc_title = title or md_path.stem
    title_p = doc.add_paragraph(doc_title)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.runs[0].bold = True
    title_p.runs[0].font.size = Pt(18)

    date_p = doc.add_paragraph(datetime.now().strftime('%Y-%m-%d %H:%M'))
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('')

    in_code = False
    code_buffer: list[str] = []

    for raw in lines:
        line = raw.rstrip('\n')

        # code fences
        if line.strip().startswith('```'):
            if not in_code:
                in_code = True
                code_buffer = []
            else:
                in_code = False
                add_code(doc, code_buffer)
                code_buffer = []
            continue

        if in_code:
            code_buffer.append(line)
            continue

        # headings
        if line.startswith('#'):
            hashes = len(line) - len(line.lstrip('#'))
            heading_text = line.lstrip('#').strip()
            if heading_text:
                add_heading(doc, heading_text, hashes)
            continue

        # lists
        stripped = line.lstrip()
        if stripped.startswith(('- ', '* ')):
            add_bullet(doc, stripped[2:])
            continue
        if stripped[:3].isdigit() and stripped[2] == '.' and stripped[3] == ' ':
            # naive numbered list detection like "1. Item"
            add_numbered(doc, stripped[4:])
            continue

        # empty line -> paragraph break
        if stripped == '':
            doc.add_paragraph('')
            continue

        # default paragraph
        add_paragraph(doc, line)

    # flush dangling code
    if in_code and code_buffer:
        add_code(doc, code_buffer)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


def main():
    ap = argparse.ArgumentParser(description='Convert Markdown file to DOCX (simple)')
    ap.add_argument('--input', '-i', required=True, help='Path to input .md')
    ap.add_argument('--output', '-o', required=True, help='Path to output .docx')
    ap.add_argument('--title', '-t', help='Optional document title')
    args = ap.parse_args()

    in_path = Path(args.input)
    out_path = Path(args.output)

    convert_markdown_to_docx(in_path, out_path, title=args.title)
    print(f"DOCX created: {out_path}")


if __name__ == '__main__':
    main()
